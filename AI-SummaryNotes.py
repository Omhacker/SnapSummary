import streamlit as st
from transformers import pipeline
import PyPDF2
import docx

# Load summarization model
summarizer = pipeline("summarization")

# Function to summarize text
def summarize_text(text, max_length=300):
    return summarizer(text, max_length=max_length, min_length=100, do_sample=False)[0]['summary_text']

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text[:3000]  # Limit to 3000 characters

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = " ".join([para.text for para in doc.paragraphs])
    return text[:3000]  # Limit to 3000 characters

# Streamlit UI with Optimized Performance
st.title("📄 SnapSummary :- AI-Powered Smart Notes Summarizer")
st.write("Summarize text, PDFs, and DOCX files in seconds!")


# User Input
user_input = st.text_area("✏️ Enter text (Max: 8000 words)", max_chars=8000)
if st.button("🚀 Summarize Text") and user_input:
    with st.spinner("🔄 Generating Summary..."):
        st.write("### Summary:")
        st.success(summarize_text(user_input))

# File Upload (PDF/DOCX)
uploaded_file = st.file_uploader("📂 Upload a PDF or DOCX", type=["pdf", "docx"])
if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]
    extracted_text = ""

    if file_type == "pdf":
        extracted_text = extract_text_from_pdf(uploaded_file)
    elif file_type == "docx":
        extracted_text = extract_text_from_docx(uploaded_file)

    if extracted_text:
        st.text_area("📖 Extracted Text:", extracted_text, height=300)
        if st.button("📝 Summarize Document"):
            with st.spinner("⏳ Summarizing..."):
                st.success(summarize_text(extracted_text))



# import streamlit as st
# from transformers import pipeline
# import PyPDF2
# import docx

# # Load optimized summarization model
# summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# # Function to summarize text efficiently
# def summarize_text(text, max_length=300):  # Optimized max length
#     return summarizer(text, max_length=max_length, min_length=100, do_sample=False)[0]['summary_text']

# # Function to extract text from PDF (optimized)
# def extract_text_from_pdf(pdf_file):
#     reader = PyPDF2.PdfReader(pdf_file)
#     text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
#     return text[:4000]  # Adjusted for performance

# # Function to extract text from DOCX (optimized)
# def extract_text_from_docx(docx_file):
#     doc = docx.Document(docx_file)
#     text = " ".join([para.text for para in doc.paragraphs])
#     return text[:4000]  # Adjusted for performance

# # Streamlit UI with Optimized Performance
# st.title("📄 AI-Powered Smart Notes Summarizer")
# st.write("Summarize text, PDFs, and DOCX files in seconds!")

# # User Input
# user_input = st.text_area("✏️ Enter text (Max: 8000 words)", max_chars=8000)
# if st.button("🚀 Summarize Text") and user_input:
#     with st.spinner("🔄 Generating Summary..."):
#         st.write("### Summary:")
#         st.success(summarize_text(user_input))

# # File Upload (PDF/DOCX)
# uploaded_file = st.file_uploader("📂 Upload a PDF or DOCX", type=["pdf", "docx"])
# if uploaded_file:
#     file_type = uploaded_file.name.split(".")[-1]
#     extracted_text = ""

#     if file_type == "pdf":
#         extracted_text = extract_text_from_pdf(uploaded_file)
#     elif file_type == "docx":
#         extracted_text = extract_text_from_docx(uploaded_file)

#     if extracted_text:
#         st.text_area("📖 Extracted Text:", extracted_text, height=300)
#         if st.button("📝 Summarize Document"):
#             with st.spinner("⏳ Summarizing..."):
#                 st.success(summarize_text(extracted_text))
